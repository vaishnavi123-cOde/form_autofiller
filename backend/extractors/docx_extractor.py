from docx import Document

def extract_docx_text(path):

    doc = Document(path)
    lines = []

    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text
                for cell in row.cells
            )
            if row_text.strip():
                lines.append(row_text)
        lines.append("")

    return "\n".join(lines)